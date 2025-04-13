"""
Document parser for different file formats.
"""

import os
import uuid
import asyncio
import logging
from typing import Dict, List, Any, Optional, BinaryIO
import fitz  # PyMuPDF for PDF
import docx   # python-docx for Word documents
import openpyxl  # For Excel files

logger = logging.getLogger(__name__)

class DocumentParser:
    """Enhanced document parser with support for multiple file formats."""
    
    def __init__(self, max_file_size_mb: int = 25):
        """
        Initialize document parser.
        
        Args:
            max_file_size_mb: Maximum file size in MB
        """
        self.max_file_size_mb = max_file_size_mb
        
    async def parse_file(self, file_content: bytes, filename: str) -> str:
        """
        Parse file content based on file type.
        
        Args:
            file_content: Binary file content
            filename: Name of the file
            
        Returns:
            Extracted text from the file
        """
        if len(file_content) > self.max_file_size_mb * 1024 * 1024:
            raise ValueError(f"File exceeds maximum size of {self.max_file_size_mb}MB")
            
        file_extension = filename.split('.')[-1].lower() if '.' in filename else ""
        
        try:
            if file_extension == 'pdf':
                return await self._parse_pdf(file_content)
            elif file_extension in ['docx', 'doc']:
                return await self._parse_docx(file_content)
            elif file_extension in ['xlsx', 'xls']:
                return await self._parse_excel(file_content)
            elif file_extension in ['txt', 'md', 'json']:
                return file_content.decode('utf-8', errors='ignore')
            else:
                # Default to treating as text
                return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Error parsing file {filename}: {str(e)}")
            # Fall back to basic text extraction
            try:
                return file_content.decode('utf-8', errors='ignore')
            except:
                raise ValueError(f"Could not parse file {filename}. Unsupported format or corrupted file.")
    
    async def _parse_pdf(self, content: bytes) -> str:
        """
        Extract text from PDF with structure preservation.
        
        Args:
            content: PDF file content
            
        Returns:
            Extracted text with structure preserved
        """
        temp_file = f"temp_{uuid.uuid4()}.pdf"
        try:
            with open(temp_file, 'wb') as f:
                f.write(content)
                
            # Let the CPU-bound task run in a thread pool
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, self._extract_pdf_text, temp_file)
                
        finally:
            if os.path.exists(temp_file):
                os.remove(temp_file)
    
    def _extract_pdf_text(self, pdf_path: str) -> str:
        """
        Helper method to extract text from PDF.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        document = fitz.open(pdf_path)
        text_blocks = []
        
        for page_num in range(len(document)):
            page = document.load_page(page_num)
            
            # Extract text blocks with position information
            blocks = page.get_text("dict")["blocks"]
            
            # Sort blocks by vertical position to preserve reading order
            for block in blocks:
                if block.get("type") == 0:  # Text block
                    block_text = ""
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            block_text += span.get("text", "")
                        block_text += "\n"
                    
                    if block_text.strip():
                        text_blocks.append(block_text.strip())
            
            # Add page separator
            text_blocks.append("\f")  # Form feed as page separator
            
        return "\n\n".join(text_blocks)
    
    async def _parse_docx(self, content: bytes) -> str:
        """
        Extract text from DOCX with structure preservation.
        
        Args:
            content: DOCX file content
            
        Returns:
            Extracted text with structure preserved
        """
        temp_file = f"temp_{uuid.uuid4()}.docx"
        try:
            with open(temp_file, 'wb') as f:
                f.write(content)
                
            # Let the CPU-bound task run in a thread pool
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, self._extract_docx_text, temp_file)
                
        finally:
            if os.path.exists(temp_file):
                os.remove(temp_file)
    
    def _extract_docx_text(self, docx_path: str) -> str:
        """
        Helper method to extract text from DOCX.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        doc = docx.Document(docx_path)
        paragraphs = []
        
        # Extract structured content including headers
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if paragraph is a heading
                if para.style.name.startswith('Heading'):
                    # Add extra newlines around headings for better structure
                    paragraphs.append(f"\n{para.text.strip()}\n")
                else:
                    paragraphs.append(para.text.strip())
        
        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                paragraphs.append("\n" + "\n".join(table_text) + "\n")
                
        return "\n\n".join(paragraphs)
    
    async def _parse_excel(self, content: bytes) -> str:
        """
        Extract text from Excel with table structure preservation.
        
        Args:
            content: Excel file content
            
        Returns:
            Extracted text with structure preserved
        """
        temp_file = f"temp_{uuid.uuid4()}.xlsx"
        try:
            with open(temp_file, 'wb') as f:
                f.write(content)
                
            # Let the CPU-bound task run in a thread pool
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, self._extract_excel_text, temp_file)
                
        finally:
            if os.path.exists(temp_file):
                os.remove(temp_file)
    
    def _extract_excel_text(self, excel_path: str) -> str:
        """
        Helper method to extract text from Excel.
        
        Args:
            excel_path: Path to Excel file
            
        Returns:
            Extracted text
        """
        workbook = openpyxl.load_workbook(excel_path, data_only=True)
        sheet_texts = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_text = [f"Sheet: {sheet_name}"]
            
            # Get the max row and column with data
            max_row = sheet.max_row
            max_col = sheet.max_column
            
            for row in range(1, max_row + 1):
                row_values = []
                for col in range(1, max_col + 1):
                    cell_value = sheet.cell(row=row, column=col).value
                    row_values.append(str(cell_value) if cell_value is not None else "")
                
                # Only include rows with content
                if any(val.strip() for val in row_values):
                    sheet_text.append(" | ".join(row_values))
            
            sheet_texts.append("\n".join(sheet_text))
            
        return "\n\n".join(sheet_texts)